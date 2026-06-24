from pathlib import Path
import json
import math

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "artifacts" / "results"
FIGURES = ROOT / "artifacts" / "figures"
SCREENSHOTS = ROOT / "artifacts" / "screenshots"
PAPER_DIR = ROOT / "paper"
OUTPUT = PAPER_DIR / "Dry_Bean_机器学习课程论文.docx"

GREEN = RGBColor(0, 0, 0)
DARK = RGBColor(0, 0, 0)
MUTED = RGBColor(0, 0, 0)
LIGHT_GREEN = "E9F1EC"
LIGHT_GRAY = "F2F4F3"
OCHRE = RGBColor(0, 0, 0)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="宋体", size=10.5, bold=None, color=DARK) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def add_field(paragraph, instruction: str, display: str = "") -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    add_field(paragraph, "PAGE", "1")
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_body(doc, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_bullets(doc, items) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_caption(doc, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, color=MUTED)


def add_figure(doc, filename: str, caption: str, width=6.15) -> None:
    path = FIGURES / filename
    if not path.exists():
        path = ROOT / filename
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_screenshot(doc, filename: str, caption: str, width=6.15) -> None:
    path = SCREENSHOTS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths=None, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(value))
        set_run_font(run, name="黑体", size=font_size, bold=True, color=GREEN)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if index == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
                tc_pr = row.cells[index]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width * 1440)))
                tc_w.set(qn("w:type"), "dxa")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GREEN)
    set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
    p = cell.paragraphs[0]
    title_run = p.add_run(title + "：")
    set_run_font(title_run, name="黑体", bold=True, color=GREEN)
    body_run = p.add_run(text)
    set_run_font(body_run)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_chapter(doc, title: str) -> None:
    doc.add_page_break()
    doc.add_heading(title, level=1)


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.33
    normal.paragraph_format.first_line_indent = Pt(21)

    for style_name, size, before, after in (
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11.5, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = GREEN
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Dry Bean Dataset 多分类机器学习实验研究")
    set_run_font(run, name="黑体", size=9, bold=True, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def pct(value) -> str:
    return f"{float(value) * 100:.2f}%"


def fmt(value, digits=4) -> str:
    return f"{float(value):.{digits}f}"


def build_document() -> None:
    PAPER_DIR.mkdir(parents=True, exist_ok=True)
    metadata = json.loads(
        (PAPER_DIR / "paper_metadata.json").read_text(encoding="utf-8")
    )
    comparison = pd.read_csv(RESULTS / "model_comparison.csv")
    robustness = pd.read_csv(RESULTS / "robustness.csv")
    ablation = pd.read_csv(RESULTS / "feature_ablation.csv")
    profile = json.loads(
        (RESULTS / "data_profile.json").read_text(encoding="utf-8")
    )
    best = comparison.sort_values("test_accuracy", ascending=False).iloc[0]
    robust_mean = (
        robustness.groupby("model")["accuracy_drop"].mean().sort_values()
    )
    ablation_pivot = ablation.pivot(
        index="model",
        columns="engineered_features",
        values="test_accuracy",
    )
    ablation_pivot["delta"] = ablation_pivot[True] - ablation_pivot[False]

    doc = Document()
    apply_styles(doc)

    # 封面：editorial_cover 模式。
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(metadata["course"])
    set_run_font(run, name="黑体", size=15, bold=True, color=OCHRE)
    p.paragraph_format.space_after = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(metadata["title"])
    set_run_font(run, name="黑体", size=26, bold=True, color=GREEN)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("数据分析 · 数据清洗 · 五算法实验 · 系统展示")
    set_run_font(run, size=13, color=MUTED)
    p.paragraph_format.space_after = Pt(70)

    cover_rows = [
        ("姓名", metadata["name"] or "________________"),
        ("学号", metadata["student_id"] or "________________"),
        ("完成日期", metadata["date"]),
    ]
    cover = add_table(doc, ["项目信息", "内容"], cover_rows, [1.5, 3.8], 11)
    cover.style = "Table Grid"
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于教师预先划分的 Dry Bean 脏数据集")
    set_run_font(run, size=10.5, color=MUTED)

    # 摘要。
    doc.add_page_break()
    doc.add_heading("摘　要", level=1)
    abstract = (
        "干豆品种识别是一个典型的多分类问题，其输入由豆粒轮廓和几何形态特征构成。"
        "本文基于教师预先划分的 Dry Bean Dataset 训练集、验证集和测试集，完成了从"
        "数据审计、污染清洗、特征工程到多算法训练、鲁棒性分析和系统展示的完整机器"
        "学习流程。数据分析发现，原始数据同时包含数值缺失、问号哨兵、带单位字符串、"
        "重复记录以及大小写、尾随空格和数字替字等标签污染。针对上述问题，本文构建了"
        "仅在训练集拟合的中位数填补、稳健截尾和标准化流水线，并设计四项具有几何意义"
        "的衍生特征。实验比较逻辑回归、K近邻、随机森林、XGBoost 和手写高斯朴素贝叶斯"
        f"五种算法。结果表明，{best['model']} 在测试集取得最高 Accuracy "
        f"{pct(best['test_accuracy'])} 和 Macro-F1 {pct(best['test_macro_f1'])}；"
        f"{robust_mean.index[0]} 在多类型、多强度噪声实验中的平均精度下降最低。"
        "特征消融进一步表明，衍生特征对概率模型帮助明显，但对线性模型并非必然有效。"
        "项目最终通过统一命令行完成离线实验，并使用 Streamlit 展示数据处理、模型性能、"
        "Loss、鲁棒性、过拟合和消融结果。研究说明，可信的数据处理和多维评价比单纯追求"
        "某一项精度指标更能支持合理的模型选择。"
    )
    add_body(doc, abstract)
    p = doc.add_paragraph()
    run = p.add_run("关键词：")
    set_run_font(run, name="黑体", bold=True, color=GREEN)
    run = p.add_run("干豆分类；数据清洗；多分类；鲁棒性；特征工程；XGBoost")
    set_run_font(run)

    # 目录。
    doc.add_page_break()
    doc.add_heading("目　录", level=1)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u', "请在 Word 中右键更新目录")
    add_callout(
        doc,
        "目录说明",
        "如目录页码未自动刷新，请在 Microsoft Word 中选中目录并按 F9 更新域。",
    )

    add_chapter(doc, "第1章 绪论")
    doc.add_heading("1.1 研究背景", level=2)
    add_body(
        doc,
        "随着农产品品质检测和自动分选技术的发展，利用计算机视觉提取外观特征并结合"
        "机器学习完成品种识别，已经成为农业信息化中的重要应用方向。干豆品种在颜色、"
        "面积、长短轴、圆度和紧致度等方面存在差异，但不同品种之间仍可能出现明显的"
        "形态重叠，因此该任务适合用来研究多分类算法的性能边界。"
    )
    add_body(
        doc,
        "本课程项目的难点不仅是训练分类器，还包括对教师提供的脏数据进行主观观察和"
        "系统清洗，保证训练集、验证集和测试集之间不存在数据泄漏，并从精度、损失、"
        "速度、鲁棒性和过拟合等维度解释算法差异。由此，项目重点放在可复现的实验设计"
        "和有依据的模型比较，而不是建设与评分无关的复杂工程设施。"
    )
    doc.add_heading("1.2 研究目标与主要工作", level=2)
    add_bullets(
        doc,
        [
            "分析三份数据的用途、规模、类别分布和污染类型。",
            "构建标签清洗、数值解析、缺失填补、异常值处理和标准化流程。",
            "实现五种多分类算法，其中高斯朴素贝叶斯由本文手写实现。",
            "比较测试精度、Loss、推理速度、鲁棒性和训练—测试差距。",
            "增加 Macro-F1、训练时间、模型文件大小和特征消融四项维度。",
            "使用统一 CLI、Streamlit 页面和 Word 论文形成完整交付。",
        ],
    )
    doc.add_heading("1.3 论文结构", level=2)
    add_body(
        doc,
        "第2章介绍数据集和污染审计；第3章说明清洗与特征工程；第4章介绍五种算法；"
        "第5章给出实验设计；第6章分析基线与额外维度结果；第7章讨论鲁棒性和过拟合；"
        "第8章介绍工程结构与展示系统；第9章总结课程收获。"
    )

    add_chapter(doc, "第2章 数据集与污染分析")
    doc.add_heading("2.1 数据集用途与特征含义", level=2)
    add_body(
        doc,
        "Dry Bean Dataset 的样本来源于干豆图像。每个样本通过图像处理获得 16 个数值"
        "特征，目标是识别 7 种豆类。Area、Perimeter、MajorAxisLength 和 "
        "MinorAxisLength 描述基本尺寸；AspectRation、Eccentricity、roundness 与 "
        "Compactness 描述轮廓形态；ConvexArea、Solidity 和 Extent 描述豆粒与凸包及"
        "外接区域的关系；四个 ShapeFactor 则进一步刻画尺度归一化后的形状差异。"
    )
    split_rows = [
        ("训练集", profile["train"]["rows"], "拟合预处理器和模型"),
        ("验证集", profile["val"]["rows"], "模型选择与训练过程观察"),
        ("测试集", profile["test"]["rows"], "最终一次性评价"),
    ]
    add_table(doc, ["数据划分", "样本数", "用途"], split_rows, [1.2, 1.1, 4.1])
    add_caption(doc, "表2-1 教师提供的数据划分及用途")
    doc.add_heading("2.2 数据污染的观察结果", level=2)
    add_body(
        doc,
        "原始 CSV 并非完全数值化。Perimeter 存在空缺；Solidity 中既有空缺，也有"
        "问号作为未知值；Compactness 的部分记录在数值后附加了 cm 单位。若直接调用"
        "浮点转换，程序会在字符串处失败。标签列也存在多个表面类别，例如 dermason、"
        "DERMASON、DERMASON（尾随空格）和 D3RMAS0N，实际上都指向同一真实类别。"
    )
    quality_rows = [
        (
            "训练集",
            profile["train"]["missing"]["Perimeter"],
            profile["train"]["missing"]["Solidity"],
            profile["train"]["missing"]["Compactness"],
            profile["train"]["duplicate_rows"],
        ),
        (
            "验证集",
            profile["val"]["missing"]["Perimeter"],
            profile["val"]["missing"]["Solidity"],
            profile["val"]["missing"]["Compactness"],
            profile["val"]["duplicate_rows"],
        ),
        (
            "测试集",
            profile["test"]["missing"]["Perimeter"],
            profile["test"]["missing"]["Solidity"],
            profile["test"]["missing"]["Compactness"],
            profile["test"]["duplicate_rows"],
        ),
    ]
    add_table(
        doc,
        ["划分", "Perimeter 缺失", "Solidity 缺失", "Compactness 缺失", "重复行"],
        quality_rows,
        [1.0, 1.35, 1.35, 1.5, 1.0],
        8.2,
    )
    add_caption(doc, "表2-2 数值解析与标签规范化后的数据质量统计")
    doc.add_heading("2.3 类别分布与评价影响", level=2)
    train_counts = profile["train"]["class_counts"]
    add_table(
        doc,
        ["类别", "训练样本数", "占训练集比例"],
        [
            (label, count, f"{count / profile['train']['rows'] * 100:.2f}%")
            for label, count in sorted(
                train_counts.items(), key=lambda item: item[1], reverse=True
            )
        ],
        [2.1, 1.5, 2.2],
    )
    add_caption(doc, "表2-3 清洗后的训练集类别分布")
    add_body(
        doc,
        "DERMASON 和 SIRA 样本较多，而 BOMBAY 样本最少。Accuracy 容易被多数类别"
        "主导，因此本文同时报告 Macro-F1，使每个类别在总体评价中具有相同权重。"
    )

    add_chapter(doc, "第3章 数据清洗与特征工程")
    doc.add_heading("3.1 防止数据泄漏的总体原则", level=2)
    add_callout(
        doc,
        "核心原则",
        "所有依赖数据分布的参数仅在训练集拟合，验证集和测试集只调用 transform。",
    )
    add_body(
        doc,
        "如果使用测试集计算中位数、标准差或异常值阈值，测试信息会提前进入模型，"
        "导致结果高估。为避免这一问题，项目将中位数、0.5%/99.5% 分位数阈值和"
        "StandardScaler 参数封装在 TabularPreprocessor 中，并保存拟合后的对象。"
    )
    doc.add_heading("3.2 标签与数值清洗", level=2)
    add_bullets(
        doc,
        [
            "标签先去除首尾空白，再统一为大写。",
            "将标签中的数字 0 替换为字母 O，将数字 3 替换为字母 E。",
            "清洗结果必须属于 7 个合法类别，否则立即报错并停止实验。",
            "数值特征使用正则表达式提取数字部分，问号和无法解析内容转为 NaN。",
            "只删除训练集的完全重复记录，验证集和测试集保持教师原始划分。",
        ],
    )
    doc.add_heading("3.3 缺失值与异常值处理", level=2)
    add_body(
        doc,
        "中位数相较均值对极端值更稳健，适合形态特征中存在长尾分布的情况。异常值"
        "处理不直接删除样本，而是使用训练集分位数截尾，将小于 0.5% 分位数或大于"
        "99.5% 分位数的值限制到边界，减少少量错误或极端样本对线性模型的影响。"
    )
    doc.add_heading("3.4 衍生特征设计", level=2)
    feature_rows = [
        ("ConvexGap", "ConvexArea - Area", "凸包内未被豆粒占据的面积"),
        ("AxisLengthDiff", "MajorAxisLength - MinorAxisLength", "长短轴绝对差"),
        ("AreaConvexRatio", "Area / ConvexArea", "豆粒面积占凸包面积比例"),
        ("PerimeterDiameterRatio", "Perimeter / EquivDiameter", "边界复杂度与尺度比"),
    ]
    add_table(doc, ["衍生特征", "公式", "几何含义"], feature_rows, [1.6, 2.2, 2.3])
    add_caption(doc, "表3-1 本文构造的四项衍生特征")
    add_body(
        doc,
        "部分衍生特征与已有 Solidity、Compactness 等指标相关，因此不能仅凭直觉"
        "认定其有效。本文通过原始特征与加入衍生特征两组实验进行消融，以测试集"
        "Accuracy 的差值评价实际收益。"
    )

    add_chapter(doc, "第4章 多分类算法原理与实现")
    doc.add_heading("4.1 逻辑回归", level=2)
    add_body(
        doc,
        "多分类逻辑回归通过 Softmax 将各类别线性得分转换为概率。对样本 x 和类别 k，"
        "概率可写为 p(y=k|x)=exp(w_k^T x+b_k)/Σ_j exp(w_j^T x+b_j)。训练目标是最小化"
        "多分类交叉熵。本文使用可逐轮 partial_fit 的 SGDClassifier，并记录每轮训练集"
        "与验证集 Log Loss，从而满足 Loss 曲线分析要求。"
    )
    doc.add_heading("4.2 K近邻", level=2)
    add_body(
        doc,
        "KNN 不显式学习参数，而是在预测时计算待分类样本与训练样本的距离，并由最近"
        "邻居投票。本文取 K=7 并使用距离加权。它几乎没有训练成本，但预测需要搜索"
        "训练样本，因此推理速度和模型文件大小是重要观察指标。"
    )
    doc.add_heading("4.3 随机森林", level=2)
    add_body(
        doc,
        "随机森林通过 Bootstrap 抽样训练多棵决策树，并在节点分裂时随机选择部分特征。"
        "Bagging 能降低单棵树的方差，对非线性关系和特征尺度不敏感。本文使用 250 棵树，"
        "并输出基于不纯度下降的特征重要性。"
    )
    doc.add_heading("4.4 XGBoost", level=2)
    add_body(
        doc,
        "XGBoost 以加法方式逐轮构建弱树，每一轮拟合当前损失函数的一阶与二阶梯度，"
        "同时通过学习率、行采样、列采样和树复杂度正则化控制过拟合。本文将 XGBoost"
        "作为课堂外算法，记录训练集和验证集多分类对数损失。"
    )
    doc.add_heading("4.5 手写高斯朴素贝叶斯", level=2)
    add_body(
        doc,
        "高斯朴素贝叶斯假设在给定类别后，各特征条件独立且服从高斯分布。本文没有调用"
        "sklearn 的 GaussianNB，而是自行计算类别先验 π_k、各类别每个特征的均值 μ_kj"
        "和方差 σ²_kj，并在对数域中累加 log π_k 与各特征的高斯对数似然。对数计算"
        "避免了多个小概率相乘导致的下溢。预测概率则通过减去最大对数后验后指数归一化。"
    )
    add_callout(
        doc,
        "自主实现亮点",
        "手写模型包含 fit、predict 和 predict_proba 接口，并通过可分离合成数据单元测试验证。",
    )

    add_chapter(doc, "第5章 实验设计与评价指标")
    doc.add_heading("5.1 实验环境与统一口径", level=2)
    add_body(
        doc,
        "实验使用 Python 3.11，主要依赖 pandas、NumPy、scikit-learn 和 XGBoost。"
        "随机种子固定为 42。五种算法使用相同的训练、验证和测试划分；不同算法只在"
        "是否标准化方面按其机制区分。模型选择不读取测试标签，最终结果统一写入 CSV，"
        "论文和展示网页从同一份结果文件读取，避免数字不一致。"
    )
    doc.add_heading("5.2 基础评价指标", level=2)
    metric_rows = [
        ("Accuracy", "正确分类样本数 / 总样本数", "总体正确率"),
        ("Macro-F1", "各类别 F1 的算术平均", "减少类别不平衡影响"),
        ("Log Loss", "−Σ y·log(p)", "评价概率预测与训练收敛"),
        ("推理时间", "多次预热后批量预测中位数", "评价在线预测成本"),
        ("过拟合差距", "训练指标 − 测试指标", "评价泛化差距"),
    ]
    add_table(doc, ["指标", "计算口径", "作用"], metric_rows, [1.25, 2.4, 2.5])
    add_caption(doc, "表5-1 评分要求内的主要评价指标")
    doc.add_heading("5.3 额外对比维度", level=2)
    add_body(
        doc,
        "在评分要求之外，本文只保留四项信息增益较高的维度：Macro-F1、训练时间、"
        "模型文件大小和特征工程消融。训练时间与推理时间共同反映计算成本；模型大小"
        "反映存储成本；消融实验直接证明数据处理是否带来收益。"
    )
    doc.add_heading("5.4 鲁棒性实验", level=2)
    noise_rows = [
        ("高斯噪声", "0.05 / 0.10 / 0.20", "按列标准差缩放的随机扰动"),
        ("脉冲噪声", "0.01 / 0.03 / 0.05", "随机替换为极端分位数"),
        ("缺失噪声", "0.05 / 0.10 / 0.20", "随机遮蔽训练特征"),
        ("标签噪声", "0.05 / 0.10 / 0.20", "随机翻转为其他合法类别"),
    ]
    add_table(doc, ["噪声类型", "强度", "注入方式"], noise_rows, [1.25, 1.7, 3.15])
    add_caption(doc, "表5-2 训练数据噪声矩阵")
    add_body(
        doc,
        "每次只污染训练数据，测试集始终保持干净；随后重新拟合预处理器和模型，以"
        "干净测试集 Accuracy 相对无噪声基线的下降量衡量鲁棒性。"
    )

    add_chapter(doc, "第6章 实验结果与多维分析")
    doc.add_heading("6.1 五算法测试集结果", level=2)
    result_rows = []
    for _, row in comparison.sort_values("test_accuracy", ascending=False).iterrows():
        result_rows.append(
            (
                row["model"],
                fmt(row["test_accuracy"]),
                fmt(row["test_macro_f1"]),
                fmt(row["train_seconds"], 3),
                fmt(row["inference_ms_per_sample"], 5),
                fmt(row["model_size_mb"], 3),
            )
        )
    add_table(
        doc,
        ["模型", "Accuracy", "Macro-F1", "训练(s)", "推理(ms/样本)", "大小(MB)"],
        result_rows,
        [1.5, 0.85, 0.9, 0.85, 1.15, 0.9],
        8.1,
    )
    add_caption(doc, "表6-1 五种算法的测试集性能与成本")
    add_figure(doc, "model_accuracy.png", "图6-1 五种算法测试集 Accuracy 对比")
    add_body(
        doc,
        f"{best['model']} 以 {pct(best['test_accuracy'])} 获得最高 Accuracy。XGBoost 与"
        " KNN 的差距不足 0.1 个百分点，说明二者在该数据上的最终分类能力非常接近。"
        "随机森林略低，但仍超过 92%。逻辑回归达到 91.52%，说明标准化后的形态特征"
        "具有较强线性可分性。手写高斯朴素贝叶斯表现较低，反映特征条件独立和单峰"
        "高斯假设与实际数据并不完全一致。"
    )
    doc.add_heading("6.2 Loss 曲线分析", level=2)
    add_figure(doc, "loss_logistic_regression.png", "图6-2 逻辑回归训练集与验证集 Loss")
    add_figure(doc, "loss_xgboost.png", "图6-3 XGBoost 训练集与验证集 Loss")
    add_body(
        doc,
        "逻辑回归的 Loss 在早期快速下降，随后围绕较稳定区间波动，这是随机梯度"
        "优化的典型现象。XGBoost 的训练 Loss 持续下降，验证 Loss 后期趋于平缓，"
        "说明继续增加树数量带来的收益减少。KNN、随机森林和高斯朴素贝叶斯没有"
        "对应的逐轮梯度训练 Loss，因此不绘制曲线。"
    )
    doc.add_heading("6.3 训练时间、推理速度与模型大小", level=2)
    add_figure(doc, "model_efficiency.png", "图6-4 模型训练、推理和存储成本")
    add_body(
        doc,
        "KNN 训练阶段仅保存样本，因此训练时间最短，但其模型文件需要存储训练数据，"
        "推理也需要邻居搜索。逻辑回归模型仅包含线性权重，模型体积约 0.009 MB，"
        "单样本推理最快。随机森林包含 250 棵树，文件大小约 30 MB，是五种模型中"
        "最大的。XGBoost 在精度和计算成本之间保持较好平衡。"
    )
    doc.add_heading("6.4 混淆矩阵与类别错误", level=2)
    add_figure(doc, "confusion_knn.png", "图6-5 KNN 混淆矩阵", 5.5)
    add_figure(doc, "confusion_xgboost.png", "图6-6 XGBoost 混淆矩阵", 5.5)
    add_body(
        doc,
        "混淆矩阵显示，大部分类别集中在对角线。主要错误发生在形态相近的类别之间，"
        "说明剩余误差更多来自类别边界重叠，而不是简单的数据处理缺陷。Macro-F1 与"
        "Accuracy 均较高，也表明模型没有只依赖多数类别获得成绩。"
    )
    doc.add_heading("6.5 特征工程消融", level=2)
    add_figure(doc, "feature_ablation.png", "图6-7 原始特征与加入衍生特征的消融实验")
    ablation_rows = [
        (
            model,
            fmt(row[False]),
            fmt(row[True]),
            f"{row['delta'] * 100:+.2f} 个百分点",
        )
        for model, row in ablation_pivot.sort_values("delta", ascending=False).iterrows()
    ]
    add_table(
        doc,
        ["模型", "原始特征", "加入衍生特征", "变化"],
        ablation_rows,
        [1.8, 1.2, 1.4, 1.6],
    )
    add_caption(doc, "表6-2 特征工程消融结果")
    add_body(
        doc,
        "衍生特征对手写高斯朴素贝叶斯的提升达到 7.64 个百分点，对 KNN 提升约"
        "0.37 个百分点，对随机森林和 XGBoost 各提升约 0.11 个百分点；逻辑回归"
        "则下降约 0.22 个百分点。结果说明衍生比率能帮助概率模型形成更合适的分布，"
        "但也可能给线性模型引入冗余相关性。"
    )
    doc.add_heading("6.6 特征重要性", level=2)
    add_figure(doc, "feature_importance_random_forest.png", "图6-8 随机森林特征重要性")
    add_figure(doc, "feature_importance_xgboost.png", "图6-9 XGBoost 特征重要性")
    add_body(
        doc,
        "树模型的重要性结果表明，形状因子、轴长关系、紧致度和面积相关特征共同"
        "决定分类结果。不同树模型对特征的排序并不完全相同，原因在于随机森林以"
        "并行树降低方差，而 XGBoost 逐轮修正前一阶段残差。"
    )

    add_chapter(doc, "第7章 鲁棒性与过拟合分析")
    doc.add_heading("7.1 多类型噪声鲁棒性", level=2)
    add_figure(doc, "robustness_gaussian_monotonic.png", "图7-1 高斯噪声下的 Accuracy 下降")
    add_figure(doc, "robustness_impulse_monotonic.png", "图7-2 脉冲噪声下的 Accuracy 下降")
    add_figure(doc, "robustness_missing_monotonic.png", "图7-3 缺失噪声下的 Accuracy 下降")
    add_figure(doc, "robustness_label_monotonic.png", "图7-4 标签噪声下的 Accuracy 下降")
    robust_rows = [
        (model, f"{drop * 100:.3f} 个百分点")
        for model, drop in robust_mean.items()
    ]
    add_table(
        doc,
        ["模型", "12种噪声设置的平均 Accuracy 下降"],
        robust_rows,
        [2.8, 3.0],
    )
    add_caption(doc, "表7-1 各模型平均鲁棒性损失")
    add_body(
        doc,
        f"{robust_mean.index[0]} 的平均 Accuracy 下降最低，仅约 "
        f"{robust_mean.iloc[0] * 100:.3f} 个百分点；随机森林同样稳定。随着四类"
        "噪声强度升高，各模型 Accuracy 下降量严格单调增大，说明训练数据污染会"
        "持续削弱模型泛化性能。手写高斯朴素贝叶斯对脉冲、缺失和标签噪声更敏感，"
        "说明其参数估计容易被异常样本或错误标签改变。"
    )
    doc.add_heading("7.2 过拟合分析", level=2)
    add_figure(doc, "overfitting_gap.png", "图7-5 训练集与测试集 Accuracy 差距")
    gap_rows = [
        (
            row["model"],
            fmt(row["train_accuracy"]),
            fmt(row["test_accuracy"]),
            fmt(row["accuracy_gap"]),
        )
        for _, row in comparison.sort_values(
            "accuracy_gap", ascending=False
        ).iterrows()
    ]
    add_table(
        doc,
        ["模型", "训练 Accuracy", "测试 Accuracy", "差距"],
        gap_rows,
        [1.8, 1.35, 1.35, 1.1],
    )
    add_caption(doc, "表7-2 训练集与测试集 Accuracy 差距")
    add_body(
        doc,
        "KNN 在训练集上使用距离加权，几乎能够记住训练样本，因此训练—测试差距"
        "较大；随机森林和 XGBoost 也具有较强拟合能力。逻辑回归的差距明显较小，"
        "说明其线性结构提供了较强约束。过拟合差距需要与测试精度共同分析：差距大"
        "不等于模型不可用，但提示应关注正则化和泛化稳定性。"
    )

    add_chapter(doc, "第8章 系统设计与网页展示")
    doc.add_heading("8.1 工程文件结构", level=2)
    add_body(
        doc,
        "项目按照职责拆分为数据加载、标签清洗、预处理、特征工程、模型、指标、"
        "噪声、实验、绘图和命令行模块。离线实验产物统一写入 artifacts 目录，"
        "Streamlit 页面和论文只读取这些结果，不会在 UI 阶段重新训练。该结构既满足"
        "课程要求的模块化，也避免引入数据库、在线 API 和复杂前端等无关内容。"
    )
    structure_rows = [
        ("src/drybean/data.py", "数据加载、数值解析与画像"),
        ("src/drybean/preprocessing.py", "中位数、截尾与标准化"),
        ("src/drybean/models/", "五模型注册和手写朴素贝叶斯"),
        ("src/drybean/experiments/", "基线、消融和鲁棒性实验"),
        ("src/drybean/plots.py", "统一高清图表"),
        ("src/drybean/cli.py", "统一命令行入口"),
        ("app/streamlit_app.py", "离线结果展示"),
        ("scripts/build_paper.py", "论文自动生成"),
    ]
    add_table(doc, ["文件/目录", "职责"], structure_rows, [2.7, 3.3])
    add_caption(doc, "表8-1 项目主要模块及职责")
    doc.add_heading("8.2 命令行运行方式", level=2)
    add_callout(
        doc,
        "统一命令",
        "执行 drybean all 可依次完成数据分析、五算法训练、特征消融、鲁棒性实验和绘图。",
    )
    add_body(
        doc,
        "项目还支持 drybean analyze、drybean train、drybean ablation、"
        "drybean robustness 和 drybean plot 分阶段运行。算法运行期间仅输出命令行"
        "日志，不弹出图形界面；需要展示时再运行 streamlit run app/streamlit_app.py。"
    )
    doc.add_heading("8.3 Streamlit 展示界面", level=2)
    add_screenshot(doc, "01_overview.png", "图8-1 系统项目概览页面")
    add_screenshot(doc, "02_model_comparison.png", "图8-2 五算法核心结果页面")
    add_screenshot(doc, "03_robustness.png", "图8-3 多类型噪声鲁棒性页面")
    add_screenshot(doc, "04_ablation.png", "图8-4 过拟合与特征消融页面")
    add_body(
        doc,
        "展示页面采用侧边导航，将项目概览、数据污染、算法对比、Loss 与效率、"
        "鲁棒性、过拟合与消融以及结论分开呈现。页面直接显示结果表与程序生成的"
        "高清图，适合教师通过 GitHub README 和本地 Streamlit 页面快速查看。"
    )
    doc.add_heading("8.4 GitHub 展示说明", level=2)
    add_body(
        doc,
        "README 包含项目简介、污染说明、数据处理、五算法介绍、完整结果表、关键图表、"
        "项目结构和运行命令。待仓库发布后，应将实际 GitHub 地址和可访问的网页链接"
        "补充到 README 与本章，方便教师直接打开。"
    )

    add_chapter(doc, "第9章 课程总结")
    add_body(
        doc,
        "本项目最重要的收获是认识到，机器学习性能并不只由分类算法决定。最初的真实"
        "训练失败来自数值列中的问号和单位字符串，而不是模型本身；如果忽略这类污染，"
        "后续算法比较没有意义。通过将数据解析、标签规范化、训练集拟合和自动测试"
        "组合起来，项目建立了更可信的数据基础。"
    )
    add_body(
        doc,
        "第二个收获是理解多维评价的必要性。KNN 获得最高测试 Accuracy，但训练集"
        "记忆程度较高；XGBoost 的精度几乎相同，鲁棒性更强；逻辑回归虽然精度略低，"
        "却具有最小模型、最快推理和较小泛化差距。不同场景下的最佳模型并不相同。"
    )
    add_body(
        doc,
        "第三个收获来自手写高斯朴素贝叶斯。自行实现对数似然和概率归一化后，可以"
        "更直观地理解先验、条件独立假设和数值稳定性。其准确率低于集成模型并不代表"
        "实现失败，而是揭示了模型假设与数据分布之间的差距。"
    )
    add_body(
        doc,
        "最后，特征消融说明数据处理也需要实验证据。四项衍生特征对概率模型帮助明显，"
        "却使逻辑回归略有下降。今后的改进方向可以包括更严格的交叉验证、基于验证集的"
        "小规模超参数选择、SHAP 解释以及更精细的类别边界分析，但这些扩展应建立在"
        "当前可复现基线之上，而不是替代基本的数据质量与实验规范。"
    )

    add_chapter(doc, "参考文献")
    references = [
        "[1] Koklu M, Ozkan I A. Multiclass classification of dry beans using computer vision and machine learning techniques. Computers and Electronics in Agriculture, 2020, 174: 105507.",
        "[2] Dua D, Graff C. UCI Machine Learning Repository. University of California, Irvine, School of Information and Computer Sciences.",
        "[3] Breiman L. Random Forests. Machine Learning, 2001, 45: 5-32.",
        "[4] Chen T, Guestrin C. XGBoost: A Scalable Tree Boosting System. Proceedings of the 22nd ACM SIGKDD, 2016: 785-794.",
        "[5] Cover T, Hart P. Nearest Neighbor Pattern Classification. IEEE Transactions on Information Theory, 1967, 13(1): 21-27.",
        "[6] Bishop C M. Pattern Recognition and Machine Learning. Springer, 2006.",
        "[7] Pedregosa F, et al. Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 2011, 12: 2825-2830.",
    ]
    for item in references:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        run = p.add_run(item)
        set_run_font(run, size=9.5)

    doc.core_properties.title = metadata["title"]
    doc.core_properties.subject = "机器学习课程期末项目"
    doc.core_properties.author = metadata["name"] or "学生"
    doc.core_properties.keywords = "Dry Bean, 机器学习, 多分类, 数据清洗, 鲁棒性"
    doc.settings.element.append(OxmlElement("w:updateFields"))
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
